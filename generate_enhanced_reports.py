"""
Generate Enhanced Reports (Markdown → HTML → PDF → DOCX)
---------------------------------------------------------
* Consolidates all earlier logic in a single, reusable script.
* Adds:
  • Automatic Table‑of‑Contents (TOC)
  • Page numbers & generous margins in PDF (WeasyPrint @page rules)
  • Clean typographic CSS for HTML/PDF
  • Optionally converts HTML → DOCX via **pypandoc** (fallback to *python‑docx* if Pandoc unavailable)
  • Graceful library‑missing checks with helpful instructions.

Usage (from project root, after notebook has written its CSV/PNG outputs):
    python generate_enhanced_reports.py
Required Python packages:
    pip install pandas openpyxl markdown weasyprint python-docx pypandoc beautifulsoup4
Note:  *pypandoc* needs Pandoc installed & on PATH – https://pandoc.org/install.html
"""

from __future__ import annotations
import os
import sys
import base64
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

# ── Helpers ────────────────────────────────────────────────────────────────
BASE_OUTPUT_DIR = Path("outputs/jupyter_notebook")
PLOTS_DIR = BASE_OUTPUT_DIR / "plots"
TABLES_DIR = BASE_OUTPUT_DIR / "tables"
REPORT_DIR = BASE_OUTPUT_DIR / "comprehensive_results"


def opath(filename: str, otype: str = "plots") -> Path:
    """Uniform path builder for notebook‑generated artefacts."""
    root = {"plots": PLOTS_DIR, "tables": TABLES_DIR, "reports": REPORT_DIR}[otype]
    return root / filename


def img2b64(path: Path) -> str:
    """Return base64‑encoded contents of *path* or empty string if missing."""
    if not path.exists():
        print(f"⚠️  Missing image: {path}")
        return ""
    return base64.b64encode(path.read_bytes()).decode()


# ── Main ───────────────────────────────────────────────────────────────────


def main() -> None:
    # 1. Dependency check ----------------------------------------------------
    missing = []
    try:
        import pandas as pd
    except ImportError:
        missing.append("pandas")
    try:
        from weasyprint import HTML, CSS  # noqa: F401 – checked import
    except ImportError:
        missing.append("weasyprint")
    try:
        import markdown  # noqa: F401
    except ImportError:
        missing.append("markdown")
    try:
        import docx  # python‑docx – only required for docx fallback
    except ImportError:
        docx = None  # type: ignore
        missing.append("python-docx")
    try:
        import pypandoc  # noqa: F401 – optional
    except ImportError:
        pypandoc = None  # type: ignore
    # Fail fast if BeautifulSoup missing – used by python‑docx fallback
    try:
        from bs4 import BeautifulSoup  # noqa: F401
    except ImportError:
        BeautifulSoup = None  # type: ignore
        missing.append("beautifulsoup4")

    if missing:
        print("❌ Missing packages: " + ", ".join(missing))
        print("Run:  pip install " + " ".join(missing))
        sys.exit(1)

    import pandas as pd  # guaranteed available after check
    from weasyprint import HTML, CSS
    import markdown

    # 2. Prepare output dirs -------------------------------------------------
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    # 3. Load notebook‑outputs ----------------------------------------------
    REQUIRED: Dict[str, str] = {
        "demo_stats": "Demographic_Clinical_Stats.csv",
        "calibrator_df": "Calibration_Table.csv",
        "h_vs_p": "H_vs_P_Results.csv",
        "g_vs_p": "G_vs_P_Results.csv",
        "h_vs_g": "H_vs_G_Results.csv",
        "model_df": "Model_Performance_Metrics.csv",
    }
    loaded: Dict[str, Any] = {}
    missing_files = []
    for key, fname in REQUIRED.items():
        fp = opath(fname, "tables")
        if fp.exists():
            idx = 0 if key == "calibrator_df" else None
            loaded[key] = pd.read_csv(fp, index_col=idx)
        else:
            missing_files.append(fname)
    if missing_files:
        print("❌ Could not locate expected file(s): " + ", ".join(missing_files))
        print("Please run the analysis notebook before generating reports.")
        sys.exit(1)

    demo_stats = loaded["demo_stats"]
    calibrator_df = loaded["calibrator_df"]
    h_vs_p = loaded["h_vs_p"]
    g_vs_p = loaded["g_vs_p"]
    h_vs_g = loaded["h_vs_g"]
    model_df = loaded["model_df"]

    sig_combined = pd.concat(
        [
            h_vs_p[(h_vs_p["FDR_Significant"]) & (h_vs_p["Log2FC_Threshold"])],
            g_vs_p[(g_vs_p["FDR_Significant"]) & (g_vs_p["Log2FC_Threshold"])],
            h_vs_g[(h_vs_g["FDR_Significant"]) & (h_vs_g["Log2FC_Threshold"])],
        ]
    ).sort_values("Q_Value")

    # 4. Build Markdown ------------------------------------------------------
    now = datetime.now().strftime("%B %d, %Y")
    md_parts = [
        "# Comprehensive miRNA Periodontal Disease Analysis Results",
        f"**Date:** {now}  ",
        "[[TOC]]",
        "\n## 1. Demographic & Clinical Characteristics",
        demo_stats.to_markdown(index=False),
        "\n## 2. ΔΔCt Transformation & QC",
        "### Calibrator Values (Healthy Group Mean ΔCt)",
        calibrator_df.to_markdown(),
    ]

    for sec, fn in [
        ("Volcano Plots", "Volcano_Plots.png"),
        ("ROC Curves", "ROC_Curves.png"),
        ("Dimensionality Reduction Plots", "Dimensionality_Reduction.png"),
    ]:
        b64 = img2b64(opath(fn, "plots"))
        if b64:
            md_parts.extend(
                [
                    f"\n## {len(md_parts)}. {sec}",
                    f"![{sec}](data:image/png;base64,{b64})",
                ]
            )

    if not sig_combined.empty:
        md_parts.extend(
            [
                "\n## Significant Differential Expression Results",
                sig_combined[
                    ["miRNA", "Comparison", "Log2_FC", "Q_Value", "Effect_Size"]
                ].to_markdown(index=False, floatfmt=".3f"),
            ]
        )

    md_content = "\n".join(md_parts)
    md_path = opath("Enhanced_Report.md", "reports")
    md_path.write_text(md_content, encoding="utf-8")
    print("✅ Markdown saved →", md_path)

    # 5. Markdown → HTML -----------------------------------------------------
    css = CSS(
        string="""
@page{size:A4;margin:2cm;@bottom-right{content:'Page ' counter(page) ' / ' counter(pages);font-size:9pt;color:#666}}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Helvetica,Arial,sans-serif;line-height:1.6}
h1,h2,h3,h4,h5{color:#003366;margin-top:1.2em}
table{border-collapse:collapse;margin:1em 0;width:auto}
th,td{border:1px solid #ddd;padding:6px 8px;font-size:10pt}
th{background:#f2f2f2}
img{max-width:100%;height:auto;border:1px solid #ddd;border-radius:4px;padding:4px}
"""
    )
    # Use proper extension names and configs for TOC
    html_body = markdown.markdown(
        md_content,
        extensions=["toc", "tables", "fenced_code"],
        extension_configs={"toc": {"permalink": True}},
    )
    html_full = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>miRNA Analysis Results</title></head><body>{html_body}</body></html>"
    html_path = opath("Enhanced_Report.html", "reports")
    html_path.write_text(html_full, encoding="utf-8")
    print("✅ HTML saved →", html_path)

    # 6. HTML → PDF ----------------------------------------------------------
    pdf_path = opath("Enhanced_Report.pdf", "reports")
    HTML(string=html_full, base_url=str(BASE_OUTPUT_DIR)).write_pdf(
        pdf_path, stylesheets=[css]
    )
    print("✅ PDF saved  →", pdf_path)

    # 7. HTML → DOCX ---------------------------------------------------------
    docx_path = opath("Enhanced_Report.docx", "reports")
    if "pypandoc" in sys.modules and sys.modules["pypandoc"]:
        import pypandoc  # type: ignore

        try:
            pypandoc.convert_text(
                html_full, "docx", format="html", outputfile=str(docx_path)
            )
            print("✅ DOCX (Pandoc) saved  →", docx_path)
        except RuntimeError as e:
            print("⚠️  Pandoc conversion failed:", e)
            sys.modules["pypandoc"] = None  # mark unusable

    if (
        (sys.modules.get("pypandoc") is None)
        and docx is not None
        and BeautifulSoup is not None
    ):
        from bs4 import BeautifulSoup  # type: ignore – already verified

        soup = BeautifulSoup(html_full, "html.parser")
        d = docx.Document()  # type: ignore
        from docx.shared import Inches  # type: ignore

        for node in soup.body.children:
            if getattr(node, "name", None) and node.name.startswith("h"):
                d.add_heading(node.get_text(), level=int(node.name[1]))
            elif getattr(node, "name", None) == "p":
                d.add_paragraph(node.get_text())
            elif getattr(node, "name", None) == "img":
                img_data = base64.b64decode(node["src"].split(",", 1)[1])
                tmp = REPORT_DIR / "_tmp.png"
                tmp.write_bytes(img_data)
                d.add_picture(str(tmp), width=Inches(6))
                tmp.unlink(missing_ok=True)
            elif getattr(node, "name", None) == "table":
                rows = node.find_all("tr")
                tbl = d.add_table(
                    rows=len(rows), cols=len(rows[0].find_all(["td", "th"]))
                )
                for r, tr in enumerate(rows):
                    for c, cell in enumerate(tr.find_all(["td", "th"])):
                        tbl.rows[r].cells[c].text = cell.get_text()
        d.save(str(docx_path))
        print("✅ DOCX (python-docx) saved →", docx_path)
    elif (sys.modules.get("pypandoc") is None) and BeautifulSoup is None:
        print(
            "⚠️  Skipped python-docx fallback because beautifulsoup4 is not installed."
        )

    print("\n🎉 All report formats generated successfully!")


if __name__ == "__main__":
    main()
